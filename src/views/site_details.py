import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime, timedelta
from src.utils.database import execute_query, get_art_forms
from src.utils.unsplash import get_site_images
from src.utils.llm import generate_site_story, generate_user_custom_site_story
import docx
from docx.shared import Inches
from fpdf import FPDF
import io

def render_site_details():
    # Get the selected site from session state
    site_name = st.session_state.get('selected_site')
    if not site_name:
        st.error("No site selected. Please select a site from the trending section.")
        if st.button("Back to Home"):
            st.session_state['current_view'] = 'home'
            st.rerun()
        return

    # Get site details from database
    site_query = """
    SELECT
        site_id,
        name,
        description,
        location,
        latitude,
        longitude,
        state,
        city,
        established_year,
        heritage_type,
        unesco_status,
        risk_level,
        health_index,
        COALESCE(story, '') as story
    FROM HERITAGE_SITES
    WHERE name = %s
    """
    site_details = execute_query(site_query, (site_name,))

    if not site_details:
        st.error("Site details not found.")
        return

    # Convert tuple to dictionary for easier access
    site = {
        'site_id': site_details[0][0],
        'name': site_details[0][1],
        'description': site_details[0][2],
        'location': site_details[0][3],
        'latitude': site_details[0][4],
        'longitude': site_details[0][5],
        'state': site_details[0][6],
        'city': site_details[0][7],
        'year_built': site_details[0][8],
        'type': site_details[0][9],
        'unesco_status': site_details[0][10],
        'risk_level': site_details[0][11],
        'health_index': site_details[0][12],
        'story': site_details[0][13]
    }

    # Get associated art forms
    art_forms = get_art_forms(site['site_id'])

    # Back button
    if st.button("← Back to Home"):
        st.session_state['current_view'] = 'home'
        st.rerun()

    # Site Header with Banner Image
    st.title(site['name'])

    # Create tabs
    tab1, tab2 = st.tabs(["Details", "Create Custom Story"])

    with tab1:
        # Get images from Unsplash
        images = get_site_images(site['name'], count=7)
        if images:
            # First row: Story and Image Grid
            col1, col2 = st.columns([4, 3])

            with col1:
                story_placeholder = st.empty()

                # Check if story exists and is valid
                story_exists = site['story'] is not None and str(site['story']).strip() != ''

                if story_exists:
                    # Display existing story
                    story_placeholder.markdown(str(site['story']))

                    # Add re-generate button
                    if st.button("🔄 Re-generate Story"):
                        story_text = ""
                        for chunk in generate_site_story(site):
                            story_text += chunk
                            story_placeholder.markdown(story_text)

                        # Save the generated story to database
                        update_query = """
                        UPDATE HERITAGE_SITES
                        SET story = %s
                        WHERE site_id = %s
                        """
                        execute_query(update_query, (story_text, site['site_id']))
                        site['story'] = story_text
                        st.rerun()

            with col2:
                # Create a 2x3 grid of images
                for i in range(0, 6, 2):
                    cols = st.columns(2)
                    for j in range(2):
                        with cols[j]:
                            st.image(images[i + j], use_container_width=True)

            st.markdown("---")

            # Second row: About and Single Image
            col1, col2 = st.columns([3, 2])

            with col1:
                st.header("About")
                st.markdown(f"**Location:** {site['location']}, {site['state']}")
                st.markdown(f"**Type:** {site['type']}")
                st.markdown(f"**Year Built:** {site['year_built']}")
                st.markdown(f"**Description:** {site['description']}")
                st.markdown(f"**UNESCO Status:** {'UNESCO World Heritage Site' if site['unesco_status'] else 'Not a UNESCO Site'}")
                st.markdown(f"**Risk Level:** {site['risk_level']}")
                st.markdown(f"**Health Index:** {site['health_index']:.2f}")

                # Display associated art forms
                if art_forms:
                    st.markdown("**Associated Art Forms:**")
                    for art_form in art_forms:
                        st.markdown(f"- **{art_form['name']}**: {art_form['description']}")

            with col2:
                st.image(images[6], use_container_width=True)

        # Rest of the existing content for the Details tab
        st.markdown("---")
        # Get visitor statistics
        stats_query = """
        SELECT * FROM VISITOR_STATS
        WHERE site_id = %s
        ORDER BY VISIT_DATE DESC
        """
        visitor_stats = execute_query(stats_query, (site['site_id'],))

        if visitor_stats:
            df_stats = pd.DataFrame(visitor_stats, columns=[
                'id', 'site_id', 'visit_date', 'visitors', 'revenue', 'season',
                'created_at', 'updated_at'
            ])
            df_stats['visit_date'] = pd.to_datetime(df_stats['visit_date'])

            # Fourth row: Daily visitors count
            fig_visitors = px.line(df_stats, x='visit_date', y='visitors',
                                 title='Daily Visitor Count')
            st.plotly_chart(fig_visitors, use_container_width=True)

            # Fifth row: Daily revenue
            fig_revenue = px.line(df_stats, x='visit_date', y='revenue',
                                title='Daily Revenue')
            st.plotly_chart(fig_revenue, use_container_width=True)

        # Sixth row: Visitor Reviews and Interactions
        st.header("Visitor Reviews & Interactions")
        try:
            # Get user interactions
            interactions_query = """
            SELECT * FROM USER_INTERACTIONS
            WHERE site_id = %s
            ORDER BY interaction_date DESC
            """
            interactions = execute_query(interactions_query, (site['site_id'],))

            if interactions:
                df_interactions = pd.DataFrame(interactions, columns=[
                    'id', 'user_id', 'site_id', 'interaction_type', 'interaction_date',
                    'rating', 'review', 'created_at', 'updated_at'
                ])

                # Display interactions in rows of 3
                for i in range(0, len(df_interactions), 3):
                    cols = st.columns(3)
                    row_reviews = df_interactions.iloc[i:i+3]
                    for idx, (_, interaction) in enumerate(row_reviews.iterrows()):
                        with cols[idx]:
                            st.markdown(f"""
                            <div style='background-color: #f0f2f6; padding: 1rem; border-radius: 0.5rem; margin: 0.5rem 0; height: 100%;'>
                                <h4 style='margin: 0 0 0.5rem 0;'>{interaction['user_id']}</h4>
                                <p style='color: #666; margin: 0 0 0.5rem 0; font-size: 0.9rem;'>{interaction['interaction_date']}</p>
                                <p style='margin: 0 0 0.5rem 0;'><strong>{interaction['interaction_type']}</strong> • {'⭐' * int(interaction['rating'])}</p>
                                <p style='margin: 0;'>{interaction['review']}</p>
                            </div>
                            """, unsafe_allow_html=True)
            else:
                st.info("No visitor interactions available for this site.")
        except Exception as e:
            st.info("Visitor interactions feature is not available at the moment.")

        # Additional image galleries
        st.markdown("---")
        st.subheader("More related photos")

        # Get more images for the additional galleries
        more_images = get_site_images(site['name'], count=20)

        if more_images:
            # Display images in rows of 4
            for i in range(0, 20, 4):
                cols = st.columns(4)
                for idx, img_url in enumerate(more_images[i:i+4]):
                    with cols[idx]:
                        st.image(img_url, use_container_width=True)

    with tab2:
        # Initialize session state variables if they don't exist
        if 'story_text' not in st.session_state:
            st.session_state.story_text = ""
        if 'story_displayed' not in st.session_state:
            st.session_state.story_displayed = False

        # Textarea for user input
        user_input = st.text_area(
            "Enter additional details or information about the heritage site:",
            height=200,
            placeholder="Share your thoughts, experiences, or any additional information about this heritage site..."
        )

        st.markdown("**Note:** Your information will be added to the story wherever it is relevant. Useless or doesn't seems to align with what the site is about will be ignored.")

        # Generate button
        if st.button("Generate Story"):
            if user_input:
                # Create a placeholder for streaming output
                story_placeholder = st.empty()
                story_text = ""

                # Generate story with streaming display
                for chunk in generate_user_custom_site_story(site, user_input):
                    story_text += chunk
                    story_placeholder.markdown(story_text)

                # Store the generated story in session state
                st.session_state.story_text = story_text
                st.session_state.story_displayed = True
                st.rerun()  # Rerun to show the story in the persistent display
            else:
                st.warning("Please enter some information before generating the story.")

        # Display the story if it exists in session state
        if st.session_state.story_displayed and st.session_state.story_text:
            st.markdown("### Generated Story")
            st.markdown(st.session_state.story_text)

        # Download buttons - only show if we have a story
        if st.session_state.story_text:
            st.markdown("### Download Options")
            col1, col2 = st.columns(2)

            with col1:
                # PDF download
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)

                # Add title
                pdf.set_font("Arial", 'B', 16)
                pdf.cell(200, 10, txt=site['name'], ln=True, align='C')
                pdf.set_font("Arial", size=12)

                # Add story
                pdf.multi_cell(0, 10, txt=st.session_state.story_text)

                # Save to bytes
                pdf_bytes = pdf.output(dest='S').encode('latin-1')
                st.download_button(
                    label="Download as PDF",
                    data=pdf_bytes,
                    file_name=f"{site['name']}_story.pdf",
                    mime="application/pdf"
                )

            with col2:
                # DOCX download
                doc = docx.Document()

                # Add title
                doc.add_heading(site['name'], 0)

                # Add story
                doc.add_paragraph(st.session_state.story_text)

                # Save to bytes
                docx_bytes = io.BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)

                st.download_button(
                    label="Download as DOCX",
                    data=docx_bytes.getvalue(),
                    file_name=f"{site['name']}_story.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    render_site_details()
